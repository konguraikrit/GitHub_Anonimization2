import json
import re
import os
import sys
import unicodedata
from faker import Faker

# Initialize Faker
fake = Faker()
fake_th = Faker("th_TH")

# Path to the accumulated sensitive data list
ACCUMULATED_LIST_PATH = "sensitive_data_list.json"
EXCEL_EXTENSIONS = {".xlsx", ".xlsm"}

# -----------------------------------------------------------------------------
# DEFINITIONS OF SENSITIVE DATA
# -----------------------------------------------------------------------------

# Generic PII regex patterns (Integrated from PrivacyAnonymizer-master + your patterns)
REGEX_PATTERNS = {
    "ssn": r"\b\d{3}-\d{2}-\d{4}\b",
    "credit_card": r"\b\d{4}[- ]?\d{4}[- ]?\d{4}[- ]?\d{4}\b",
    "coordinates": r'\b\d{1,2}°\s\d{1,2}\'\s\d{1,2}\.\d"\s[NSEW]\b',
    "project_code": r'\b\d{1,2}/\d{4}\b',
    "date": r'(?:\b\d{4}-\d{2}-\d{2}\b|[0-3]?\d\s+(?:มกราคม|กุมภาพันธ์|มีนาคม|เมษายน|พฤษภาคม|มิถุนายน|กรกฎาคม|สิงหาคม|กันยายน|ตุลาคม|พฤศจิกายน|ธันวาคม)\s+[12]\d{3})',
    "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    "thai_id": r'\b\d-\d{4}-\d{5}-\d{2}-\d\b',
    "phone": r'(?<!\d)(?:\+66|0)[689]\d[- ]?\d{3}[- ]?\d{4}(?!\d)',
    "name": r'(?:นาย|นางสาว|นาง|เด็กชาย|เด็กหญิง|ดร\.|ศ\.(?:\s*ดร\.)?|รศ\.(?:\s*ดร\.)?|ผศ\.(?:\s*ดร\.)?|นพ\.|พญ\.|ทนาย|คุณ)\s*[\u0E00-\u0E7F]{2,30}(?:[ \t]+[\u0E00-\u0E7F]{2,30}){0,2}',
    "company": r'(?:การไฟฟ้าฝ่ายผลิตแห่งประเทศไทย|กฟผ\.|บริษัท\s+[\u0E00-\u0E7F\w\s]+?\s+จำกัด(?:\s+\(มหาชน\))?|ห้างหุ้นส่วน(?:\s*จำกัด)?\s+[\u0E00-\u0E7F\w]{2,40}|มูลนิธิ[\u0E00-\u0E7F\w]+(?:\s+[\u0E00-\u0E7F\w]+){0,4}|สมาคม[\u0E00-\u0E7F\w]+(?:\s+[\u0E00-\u0E7F\w]+){0,4})',
    "address": r'เลขที่\s*[0-9A-Za-z/.-]+(?:\s+หมู่\s*\d+)?(?:\s+ถนน[\u0E00-\u0E7F0-9\s]+?)?\s+(?:แขวง|ตำบล)[\u0E00-\u0E7F]+\s+(?:เขต|อำเภอ)[\u0E00-\u0E7F]+\s+(?:จังหวัด)?[\u0E00-\u0E7F]+\s*\d{5}',
    "unit_range": r'\bUnits?\s+\d+(?:-\d+)?\b',
    "doc_code": r"\b[A-Z0-9]+(?:-[A-Z0-9]+){3,}\b",
    "timestamp": r"\b\d{4}-\d{2}-\d{2}[ T]\d{2}:\d{2}:\d{2}(?:\.\d+)?(?:Z|[+-]\d{2}:?\d{2})?\b",
    "hex_encoded": r"_x[0-9a-fA-F]{4}_",
    "author_name": r'"author"\s*:\s*"([^"]+)"',
    "doc_title": r'"title"\s*:\s*"([^"]+)"',
    "line_id": r'(?:Line\s*[Ii][Dd]|LINE\s*ID|\u0e44\u0e25\u0e19\u0e4c)[\s::\uff1a]*([@\w._-]{3,30})',
    "passport": r'(?:Passport\s*(?:No\.?|Number)?|\u0e2b\u0e19\u0e31\u0e07\u0e2a\u0e37\u0e2d\u0e40\u0e14\u0e34\u0e19\u0e17\u0e32\u0e07(?:\s*\u0e40\u0e25\u0e02\u0e17\u0e35\u0e48)?)[\s:.]*([A-Z]{2}\d{7})\b',
}

def normalize_text(text):
    if text is None:
        return ""
    text = unicodedata.normalize("NFC", text)
    return re.sub(r"[\u200B-\u200D\uFEFF]", "", text)

def contains_thai(text):
    return bool(re.search(r"[\u0E00-\u0E7F]", text or ""))

def load_accumulated_list():
    if os.path.exists(ACCUMULATED_LIST_PATH):
        try:
            with open(ACCUMULATED_LIST_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Warning: Could not load accumulated list: {e}")
    return []

def save_accumulated_list(data_list):
    try:
        # Keep only unique patterns
        unique_list = []
        seen = set()
        for item in data_list:
            key = (item['pattern'], item['category'], item.get('sensitive', False))
            if key not in seen:
                unique_list.append(item)
                seen.add(key)
        
        with open(ACCUMULATED_LIST_PATH, 'w', encoding='utf-8') as f:
            json.dump(unique_list, f, indent=4, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving accumulated list: {e}")

def _excel_cell_to_json(value, pd):
    """Convert pandas/openpyxl cell values into JSON-friendly primitives."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass

    if hasattr(value, "isoformat") and not isinstance(value, str):
        try:
            return str(value)
        except (TypeError, ValueError):
            pass

    if hasattr(value, "item"):
        try:
            return value.item()
        except (TypeError, ValueError):
            pass

    return value

def convert_excel_to_genai(excel_path):
    """Convert an Excel workbook to a JSON-like dict for scanning/anonymization."""
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError(
            "Excel support requires pandas and openpyxl. "
            "Install with: pip install pandas openpyxl"
        ) from exc

    try:
        sheets_dict = pd.read_excel(
            excel_path,
            sheet_name=None,
            engine="openpyxl",
            header=None
        )
    except ImportError as exc:
        raise RuntimeError(
            "Excel support requires openpyxl for .xlsx/.xlsm files. "
            "Install with: pip install openpyxl"
        ) from exc

    combined_data = {
        "metadata": {
            "source_type": "EXCEL",
            "sheet_count": len(sheets_dict),
            "sheets": list(sheets_dict.keys())
        },
        "sheets": {}
    }

    for sheet_name, df in sheets_dict.items():
        # Match the reference converter: every row is data, no header inference.
        df = df.dropna(how="all").reset_index(drop=True)

        if df.shape[1] == 1:
            sheet_data = [
                _excel_cell_to_json(value, pd)
                for value in df.iloc[:, 0].tolist()
            ]
        else:
            sheet_data = [
                [_excel_cell_to_json(value, pd) for value in row]
                for row in df.values.tolist()
            ]

        combined_data["sheets"][str(sheet_name)] = sheet_data

    return combined_data

# -----------------------------------------------------------------------------
# ANONYMIZATION LOGIC (FAKER INTEGRATION)
# -----------------------------------------------------------------------------

REPLACEMENT_CACHE = {}

def get_fake_value(category, original_text):
    cache_key = (category, original_text.lower())
    if cache_key in REPLACEMENT_CACHE:
        return REPLACEMENT_CACHE[cache_key]

    thai_original = contains_thai(original_text)

    mapping = {
        "company": lambda: f"[Company {fake_th.company() if thai_original else fake.company()}]",
        "country": lambda: f"[Country {fake.country()}]",
        "city": lambda: f"[City {fake_th.city() if thai_original else fake.city()}]",
        "province": lambda: f"[Location {fake_th.city() if thai_original else fake.city()}]",
        "location": lambda: f"[Location {fake_th.city() if thai_original else fake.city()}]",
        "university": lambda: f"[University {fake.company()}]",
        "coordinates": lambda: f"[Lat/Long {fake.latitude()}, {fake.longitude()}]",
        "project_code": lambda: f"[Code {fake.bothify(text='##/####')}]",
        "doc_code": lambda: f"[DocCode {fake.bothify(text='???-###-?-###-###-###-?')}]",
        "date": lambda: fake.date(),
        "timestamp": lambda: f"{fake.date()} {fake.time()}+00:00",
        "ssn": lambda: fake.ssn(),
        "thai_id": lambda: fake.bothify(text="#-####-#####-##-#"),
        "phone": lambda: fake_th.phone_number() if thai_original else fake.phone_number(),
        "credit_card": lambda: fake.credit_card_number(),
        "email": lambda: fake.email(),
        "name": lambda: fake_th.name() if thai_original else fake.name(),
        "author_name": lambda: fake_th.name() if thai_original else fake.name(),
        "address": lambda: fake_th.address() if thai_original else fake.address(),
        "doc_title": lambda: f"[Title {fake.catch_phrase()}]",
        "hex_encoded": lambda: f"[Encoded {fake.hexify(text='^^^^')}]",
        "passport": lambda: fake.bothify(text="??#######").upper(),
        "line_id": lambda: f"@{fake.user_name()}",
        "money": lambda: f"{fake.numerify(text='##,###')} บาท",
        "percent": lambda: f"{fake.random_int(min=1, max=100)}%",
        "url": lambda: fake.url(),
    }
    
    gen = mapping.get(category, lambda: f"[{category.upper()} {fake.random_number(digits=4)}]")
    val = gen()
    REPLACEMENT_CACHE[cache_key] = val
    return val

# -----------------------------------------------------------------------------

# -----------------------------------------------------------------------------
# THAI NER — optional detection via PyThaiNLP (pip install pythainlp)
# Enable by setting environment variable: THAI_NER=1
# -----------------------------------------------------------------------------

NER_CATEGORY_MAP = {
    "PERSON": "name",
    "ORGANIZATION": "company",
    "LOCATION": "location",
    "DATE": "date",
    "TIME": "timestamp",
    "MONEY": "money",
    "PERCENT": "percent",
    "EMAIL": "email",
    "PHONE": "phone",
    "URL": "url",
}

_NER_TAGGER = None


def _get_ner_tagger():
    global _NER_TAGGER
    if _NER_TAGGER is None:
        from pythainlp.ner import NER
        _NER_TAGGER = NER("thainer-corpus-v2-base-crf")
    return _NER_TAGGER


def _flush_ner_entity(entity_text, ner_type, source_text, found_items):
    entity_text = entity_text.strip()
    if len(entity_text) < 2 or entity_text not in source_text:
        return
    cat = NER_CATEGORY_MAP.get(ner_type, ner_type.lower())
    if not any(item["pattern"] == entity_text for item in found_items):
        found_items.append({
            "pattern": entity_text,
            "category": cat,
            "sensitive": True,
            "type": "NER",
            "count": source_text.count(entity_text)
        })


def scan_with_thai_ner(text):
    """Run PyThaiNLP NER on a single Thai text string. Returns [] if unavailable."""
    try:
        tagger = _get_ner_tagger()
    except Exception:
        return []
    found_items = []
    try:
        tagged = tagger.get_ner(text, pos=False)
        current_tokens, current_type = [], None
        for word, tag in tagged:
            if tag.startswith("B-"):
                if current_tokens and current_type:
                    _flush_ner_entity("".join(current_tokens), current_type, text, found_items)
                current_tokens, current_type = [word], tag[2:]
            elif tag.startswith("I-") and current_type == tag[2:]:
                current_tokens.append(word)
            else:
                if current_tokens and current_type:
                    _flush_ner_entity("".join(current_tokens), current_type, text, found_items)
                current_tokens, current_type = [], None
        if current_tokens and current_type:
            _flush_ner_entity("".join(current_tokens), current_type, text, found_items)
    except Exception:
        pass
    return found_items


def collect_ner_items(data):
    """Walk a data structure or plain text string, running NER on Thai string leaves.
    Requires THAI_NER=1 env var and: pip install pythainlp"""
    if not os.environ.get("THAI_NER", ""):
        return []
    all_items = []
    seen = set()

    def _walk(node):
        if isinstance(node, str):
            if not contains_thai(node):
                return
            for item in scan_with_thai_ner(node):
                p = item["pattern"]
                if p not in seen:
                    seen.add(p)
                    all_items.append(item)
        elif isinstance(node, dict):
            for v in node.values():
                _walk(v)
        elif isinstance(node, list):
            for elem in node:
                _walk(elem)

    _walk(data)
    return all_items

# FILE HANDLING (Integrated from PrivacyAnonymizer-master)
# -----------------------------------------------------------------------------

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        try:
            import PyPDF2
        except ImportError as exc:
            raise RuntimeError("PDF support requires PyPDF2. Install with: pip install PyPDF2") from exc
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages])
    elif ext == '.docx':
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("DOCX support requires python-docx. Install with: pip install python-docx") from exc
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == '.json':
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.dumps(json.load(f), ensure_ascii=False)
    elif ext in EXCEL_EXTENSIONS:
        return json.dumps(convert_excel_to_genai(file_path), ensure_ascii=False, default=str)
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

# -----------------------------------------------------------------------------
# CORE SCANNING & REPLACEMENT
# -----------------------------------------------------------------------------

def scan_text(text, accumulated_list):
    """Scans text for both accumulated literals and generic regex patterns."""
    text = normalize_text(text)
    found_items = [] # List of dicts: {pattern, category, sensitive, type (literal/regex), count}
    
    # 1. Scan for Literals (from accumulated list)
    for item in accumulated_list:
        pattern = item['pattern']
        sensitive = item.get('sensitive', False)
        flags = 0 if sensitive else re.IGNORECASE
        found = re.findall(re.escape(pattern), text, flags)
        if found:
            # Group by exact found string (case may vary if insensitive)
            for m in set(found):
                found_items.append({
                    "pattern": m,
                    "category": item['category'],
                    "sensitive": sensitive,
                    "type": "Accumulated",
                    "count": text.count(m)
                })

    # 2. Scan for Regex Patterns
    for cat, pattern in REGEX_PATTERNS.items():
        found = re.findall(pattern, text)
        if found:
            # Group unique matches found by this regex
            unique_matches = set(found)
            for m in unique_matches:
                # Check if already caught by literals to avoid duplicates
                if not any(item['pattern'] == m for item in found_items):
                    found_items.append({
                        "pattern": m,
                        "category": cat,
                        "sensitive": True,
                        "type": "New (Regex)",
                        "count": text.count(m)
                    })
                    
    # 3. Simple Name detection (from PrivacyAnonymizer-master) - Optional
    # We add this but it might be noisy, so we label it clearly.
    # name_pattern = r"\b[A-Z][a-z]+ [A-Z][a-z]+\b" # Full Name attempt
    # ... (skipping for now to avoid noise, user can add specific names to list)

    return found_items

def prompt_for_manual_names(text):
    """Allows reviewers to add exact Thai/person names missed by automatic detection."""
    manual = input("Specific names to anonymize? Type comma-separated names, or press Enter to skip: ").strip()
    if not manual:
        return []

    text = normalize_text(text)
    manual_items = []
    for raw_name in manual.split(','):
        name = normalize_text(raw_name.strip())
        if not name:
            continue
        flags = 0 if contains_thai(name) else re.IGNORECASE
        found = re.findall(re.escape(name), text, flags)
        if found:
            manual_items.append({
                "pattern": name,
                "category": "name",
                "sensitive": True,
                "type": "Manual",
                "count": len(found)
            })
        else:
            print(f"Warning: manual name not found in text: {name}")
    return manual_items

def apply_anonymization(text, approved_items):
    """Replaces approved items in text."""
    log = []
    # Sort by length descending to prevent partial replacement bugs
    sorted_items = sorted(approved_items, key=lambda x: len(x['pattern']), reverse=True)
    
    current_text = normalize_text(text)
    for item in sorted_items:
        pattern = item['pattern']
        category = item['category']
        sensitive = item.get('sensitive', False)
        flags = 0 if sensitive else re.IGNORECASE
        
        # Check if exists (might have been replaced by a longer pattern)
        if re.search(re.escape(pattern), current_text, flags):
            replacement = get_fake_value(category, pattern)
            current_text = re.sub(re.escape(pattern), replacement, current_text, flags=flags)
            log.append({"original": pattern, "replacement": replacement, "category": category})
            
    return current_text, log

def process_json_recursively(structured_data, approved_items):
    full_log = []
    if isinstance(data, dict):
        new_dict = {}
        for k, v in data.items():
            new_val, log = process_json_recursively(v, approved_items)
            new_dict[k] = new_val
            full_log.extend(log)
        return new_dict, full_log
    elif isinstance(data, list):
        new_list = []
        for item in data:
            new_val, log = process_json_recursively(item, approved_items)
            new_list.append(new_val)
            full_log.extend(log)
        return new_list, full_log
    elif isinstance(data, str):
        return apply_anonymization(data, approved_items)
    return data, []

# -----------------------------------------------------------------------------
# MAIN CLI
# -----------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python anonymize_data_pro_v3.py <file_path>")
        print("Supported: .json, .pdf, .docx, .txt, .xlsx, .xlsm")
        return

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"Error: {file_path} not found.")
        return

    # 1. Load data
    accumulated_list = load_accumulated_list()
    ext = os.path.splitext(file_path)[1].lower()

    structured_data = None
    if ext == '.json':
        with open(file_path, 'r', encoding='utf-8') as f:
            structured_data = json.load(f)
        text_for_scanning = json.dumps(structured_data, ensure_ascii=False)
    elif ext in EXCEL_EXTENSIONS:
        structured_data = convert_excel_to_genai(file_path)
        text_for_scanning = json.dumps(structured_data, ensure_ascii=False, default=str)
    else:
        text_for_scanning = extract_text_from_file(file_path)

    # 2. Scan
    print(f"[SCAN] Scanning {file_path}...")
    found_items = scan_text(text_for_scanning, accumulated_list)

    # NER scan on structured/raw data (set THAI_NER=1 to enable)
    ner_source = structured_data if structured_data is not None else text_for_scanning
    ner_items = collect_ner_items(ner_source)
    for item in ner_items:
        if not any(existing['pattern'] == item['pattern'] for existing in found_items):
            found_items.append(item)

    if not found_items:
        print("[OK] No sensitive data detected.")
        approved_items = prompt_for_manual_names(text_for_scanning)
        if not approved_items:
            return
    else:
        # 3. Review & Approval
        print("\n--- SENSITIVE DATA DETECTED ---")
        print(f"{'ID':<4} {'Type':<15} {'Category':<15} {'Count':<6} {'Pattern/Match'}")
        print("-" * 80)
        for i, item in enumerate(found_items):
            print(f"{i+1:<4} {item['type']:<15} {item['category']:<15} {item['count']:<6} {item['pattern']}")

        print("\n[Options] 'y': Approve All, 'n': Cancel, or comma-separated IDs (e.g., 1,3,5)")
        choice = input("Your selection: ").lower().strip()

        approved_items = []
        if choice == 'y':
            approved_items = found_items
        elif choice == 'n' or not choice:
            print("Operation cancelled.")
            return
        else:
            try:
                indices = [int(x.strip()) - 1 for x in choice.split(',')]
                for idx in indices:
                    if 0 <= idx < len(found_items):
                        approved_items.append(found_items[idx])
            except ValueError:
                print("[ERROR] Invalid input.")
                return

        approved_items.extend(prompt_for_manual_names(text_for_scanning))

    if not approved_items:
        print("No items selected.")
        return

    # 4. Update Accumulated List with new Regex detections
    newly_approved_literals = []
    for item in approved_items:
        if item['type'] != "Accumulated":
            newly_approved_literals.append({
                "pattern": item['pattern'],
                "category": item['category'],
                "sensitive": item['sensitive']
            })
    
    if newly_approved_literals:
        print(f"[SAVE] Updating accumulated list with {len(newly_approved_literals)} new items...")
        save_accumulated_list(accumulated_list + newly_approved_literals)

    # 5. Execute Anonymization
    print("[ANONYMIZE] Anonymizing...")
    
    if ext == '.json':
        anonymized_data, log = process_json_recursively(structured_data, approved_items)
        output_path = file_path.replace('.json', '_pro_anonymized.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(anonymized_data, f, indent=4, ensure_ascii=False)
    elif ext in EXCEL_EXTENSIONS:
        anonymized_data, log = process_json_recursively(structured_data, approved_items)
        output_path = f"{os.path.splitext(file_path)[0]}_pro_anonymized.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(anonymized_data, f, indent=4, ensure_ascii=False, default=str)
    elif ext == '.docx':
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("DOCX support requires python-docx. Install with: pip install python-docx") from exc
        doc = docx.Document(file_path)
        log = []
        for para in doc.paragraphs:
            new_text, p_log = apply_anonymization(para.text, approved_items)
            para.text = new_text
            log.extend(p_log)
        output_path = file_path.replace('.docx', '_pro_anonymized.docx')
        doc.save(output_path)
    else:
        # Handles TXT and PDF (output as TXT for PDF as we can't easily write back to PDF structure)
        anonymized_text, log = apply_anonymization(text_for_scanning, approved_items)
        suffix = "_pro_anonymized.txt" if ext == '.pdf' else f"_pro_anonymized{ext}"
        output_path = os.path.splitext(file_path)[0] + suffix
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(anonymized_text)

    # 6. Logging
    log_file = "anonymization_pro_log.txt"
    with open(log_file, 'a', encoding='utf-8') as f:
        f.write(f"\n--- SESSION: {file_path} ---\n")
        for entry in log:
            f.write(f"[{entry['category']}] {entry['original']} -> {entry['replacement']}\n")

    print(f"[OK] Success! Saved to: {output_path}")
    print(f"[LOG] Log updated: {log_file}")

if __name__ == "__main__":
    main()
